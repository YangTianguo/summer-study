from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ==================== 全局样式设置 ====================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_font(cell, text, bold=False, size=Pt(10), color=None, font_name='宋体'):
    """设置单元格字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color

def add_heading_styled(doc, text, level=1):
    """添加格式化标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading

def add_paragraph_styled(doc, text, bold=False, size=Pt(11), indent=False):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def create_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, header, bold=True, size=Pt(10), color=RGBColor(255, 255, 255))
        set_cell_shading(cell, '2F5496')

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            set_cell_font(cell, str(val), size=Pt(10))
            if r % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

# ==================== 封面/标题 ====================
doc.add_paragraph()  # 空行
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('杨田安 暑期提升计划')
run.font.name = '黑体'
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(47, 84, 150)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('七年级（6）班 · 2026年暑假 · 七升八衔接')
run.font.name = '楷体'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(89, 89, 89)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

# 基本信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_text = '\n制订日期：2026年7月14日  |  学生：杨田安  |  七（6）班  |  校排名：175  |  班级排名：24'
run = info.add_run(info_text)
run.font.name = '宋体'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ==================== 一、成绩诊断分析 ====================
add_heading_styled(doc, '一、成绩诊断分析', level=1)

add_paragraph_styled(doc, '杨田安同学七年级下学期期末考试成绩总分234.5分，位列学校第175名、班级第24名。各科成绩及问题诊断如下：')

# 成绩表
create_table(doc,
    ['科目', '得分', '得分率', '问题诊断'],
    [
        ['语文', '68', '56.7%', '基础尚可，阅读理解和作文是主要失分点'],
        ['数学', '35', '29.2%', '⚠ 严重薄弱：基础知识漏洞大，计算能力差，需从七年级上重新梳理'],
        ['英语', '30.5', '25.4%', '⚠ 严重薄弱：词汇量严重不足，语法零基础，音标可能未掌握'],
        ['生物', '47', '~47%', '有提升空间，概念理解不够深入'],
        ['政治', '26', '~26%', '⚠ 严重薄弱：知识点未背诵，答题没有框架'],
        ['历史', '14', '~14%', '🔴 极薄弱：时间线混乱，重大事件完全未记忆'],
        ['地理', '14', '~14%', '🔴 极薄弱：基础概念空白，地图识读能力为零'],
    ],
    col_widths=[2.5, 1.5, 1.5, 9.5]
)

doc.add_paragraph()  # 空行

# 优先级表
add_paragraph_styled(doc, '薄弱科目按紧急程度排序：', bold=True)
create_table(doc,
    ['优先级', '科目', '理由'],
    [
        ['🔴 第一梯队', '数学', '初二新增物理，数学是基础；一旦掉队，物理也无法学好'],
        ['🔴 第一梯队', '英语', '语言学科靠积累，暑假不补，初二更跟不上；直接影响中考'],
        ['🟡 第二梯队', '历史 / 地理 / 政治', '背诵类科目，暑假集中记忆可快速提分，性价比极高'],
        ['🟢 第三梯队', '生物 / 语文', '有一定基础，保持巩固即可，不需要占用核心时间'],
    ],
    col_widths=[2.5, 3.5, 9]
)

# ==================== 二、总体目标 ====================
add_heading_styled(doc, '二、总体目标', level=1)

add_paragraph_styled(doc, '以初二上学期期中考试为检验节点，各科目标如下：', indent=True)

create_table(doc,
    ['科目', '当前分数', '目标分数', '提升幅度', '目标得分率'],
    [
        ['语文', '68', '80+', '+12', '67%'],
        ['数学', '35', '65+', '+30', '54%'],
        ['英语', '30.5', '60+', '+30', '50%'],
        ['生物', '47', '60+', '+13', '60%'],
        ['政治', '26', '55+', '+29', '55%'],
        ['历史', '14', '50+', '+36', '50%'],
        ['地理', '14', '50+', '+36', '50%'],
        ['总分', '234.5', '420+', '+185', '—'],
    ],
    col_widths=[2.5, 2.5, 2.5, 2.5, 3]
)

doc.add_paragraph()
add_paragraph_styled(doc, '💡 提升185分看起来很多，但分解到每科不过是把得分率从30%提到50-60%。把七年级基础补扎实，这个目标是完全可实现的。', bold=True, indent=True)

# ==================== 三、每日作息表 ====================
add_heading_styled(doc, '三、每日作息表（周一至周六）', level=1)

add_paragraph_styled(doc, '暑假共8周。每周学习6天，周日为休息日（可安排1-2小时轻松回顾）。每天学习约9-10小时，重点时段分配给数学和英语。', indent=True)

create_table(doc,
    ['时间段', '时长', '科目', '具体内容'],
    [
        ['7:00 - 7:30', '30分钟', '——', '起床 + 早餐（固定作息，不熬夜）'],
        ['7:30 - 8:00', '30分钟', '🏴 英语晨读', '大声朗读课文 / 单词（必须出声读，培养语感）'],
        ['8:00 - 9:30', '90分钟', '🔢 数学（上午）', '一天中精力最好的时段给数学——复习七年级章节'],
        ['9:30 - 9:45', '15分钟', '——', '休息（活动身体，不碰手机）'],
        ['9:45 - 11:00', '75分钟', '🏴 英语（上午）', '词汇 + 语法 + 听力训练'],
        ['11:00 - 11:15', '15分钟', '——', '休息'],
        ['11:15 - 12:00', '45分钟', '📖 小科轮换', '历史 / 地理 / 政治每日轮换一门'],
        ['12:00 - 14:00', '120分钟', '——', '午餐 + 午休（午睡30-40分钟，保持下午精力充沛）'],
        ['14:00 - 15:30', '90分钟', '📚 语文', '古诗文背诵 + 阅读理解 + 文言文'],
        ['15:30 - 15:45', '15分钟', '——', '休息'],
        ['15:45 - 16:45', '60分钟', '🧬 生物', '复习初一薄弱章节 + 预习初二内容'],
        ['16:45 - 17:00', '15分钟', '——', '休息'],
        ['17:00 - 18:00', '60分钟', '📝 复习+错题', '回顾当天所学内容 + 整理错题本'],
        ['18:00 - 19:00', '60分钟', '——', '晚餐 + 自由活动'],
        ['19:00 - 20:30', '90分钟', '🔢 数学（晚上）', '计算专项训练 + 当天错题重做'],
        ['20:30 - 21:00', '30分钟', '🏴 英语单词', '用单词卡 / APP 复习当日+之前所有单词'],
        ['21:00 - 21:30', '30分钟', '📖 自由阅读', '课外阅读 / 英语听力磨耳朵（轻松学习）'],
        ['21:30 - 22:00', '30分钟', '——', '洗漱准备就寝'],
        ['22:00', '——', '——', '熄灯睡觉（保证8小时以上睡眠）'],
    ],
    col_widths=[2.8, 1.8, 2.5, 7.5]
)

add_paragraph_styled(doc, '⚠ 周日为休息日：可安排1-2小时轻松回顾当周内容，其余时间自由安排。劳逸结合才能持久。', bold=True, indent=True)

# ==================== 四、分科学习策略 ====================
add_heading_styled(doc, '四、分科学习策略', level=1)

# --- 数学 ---
add_heading_styled(doc, '🔢 数学（每日约3小时）', level=2)
add_paragraph_styled(doc, '核心问题：七年级基础计算不过关，概念理解模糊，知识体系存在大面积空白。', bold=True, indent=True)

add_paragraph_styled(doc, '八周学习进度：', bold=True)
create_table(doc,
    ['阶段', '时间', '内容', '学习重点'],
    [
        ['第一阶段', '第1-2周', '有理数运算（七上第一章）', '正负数加减乘除、乘方、科学记数法——计算是根基'],
        ['第二阶段', '第3-4周', '整式加减 + 一元一次方程（七上第二、三章）', '从数字到字母的抽象过渡，方程思维建立'],
        ['第三阶段', '第5-6周', '二元一次方程组 + 不等式（七下）', '消元法、代入法，应用题建模'],
        ['第四阶段', '第7-8周', '几何初步（线段、角、平行线）', '为初二几何证明打基础 + 预习初二上实数'],
    ],
    col_widths=[2.5, 2, 5, 5]
)

doc.add_paragraph()
add_paragraph_styled(doc, '每日学习方法：', bold=True)
add_paragraph_styled(doc, '① 每天先做20道计算题（限时20分钟），这是基本功中的基本功，计算不过关后面全白学。', indent=True)
add_paragraph_styled(doc, '② "三步学习法"：看课本例题 → 盖住答案自己写 → 对照答案 → 找出错误原因并记录。', indent=True)
add_paragraph_styled(doc, '③ 错题本是数学提分最关键的工具：每道错题必须写清楚【错因+正确解法+同类题1道】。', indent=True)
add_paragraph_styled(doc, '④ 每周日做一套综合小测（只做基础题和中档题），检验本周学习效果。', indent=True)
add_paragraph_styled(doc, '推荐教辅：《五年中考三年模拟》七年级分册，只做基础篇和提升篇，拔高篇暂时跳过。', indent=True)

# --- 英语 ---
add_heading_styled(doc, '🏴 英语（每日约2.5小时）', level=2)
add_paragraph_styled(doc, '核心问题：词汇量极度匮乏（估计不足200词），语法概念为零，音标未掌握导致无法自主读单词。', bold=True, indent=True)

add_paragraph_styled(doc, '八周学习进度：', bold=True)
create_table(doc,
    ['阶段', '时间', '内容', '具体任务'],
    [
        ['全程', '每天', '单词滚动记忆', '每天背15个新词 + 复习之前所有单词（艾宾浩斯记忆法）'],
        ['第一阶段', '第1-3周', '音标攻克', '48个国际音标全部掌握，做到"见词能读、听音能写"'],
        ['第二阶段', '第1-4周', '七上单词+课文', '七上所有单词和重点课文过一遍，课文要求熟读背诵'],
        ['第三阶段', '第5-6周', '七下单词+语法', '一般现在时、现在进行时、人称代词、名词复数'],
        ['第四阶段', '第7-8周', '初二上预习', '一般过去时 + There be句型 + 初二上册前2单元单词预习'],
    ],
    col_widths=[2.5, 2, 3.5, 6.5]
)

doc.add_paragraph()
add_paragraph_styled(doc, '学习方法要点：', bold=True)
add_paragraph_styled(doc, '① 音标是第一关：B站搜索"英语音标教学"，每天学4-5个，一周内过完全部48个音标。学会音标后，即使不认识单词也能读出来，自学能力会质变。', indent=True)
add_paragraph_styled(doc, '② 单词记忆法——艾宾浩斯遗忘曲线：新词当天复习3次（早中晚各一次），然后在第1天、第2天、第4天、第7天、第15天各复习一次。推荐使用"百词斩"APP辅助。', indent=True)
add_paragraph_styled(doc, '③ 语法不要死记规则，每学一个语法点就自己造5个句子。比如学了现在进行时，就造"I am reading a book"这样的句子，说出来才算真会了。', indent=True)
add_paragraph_styled(doc, '④ 听力"磨耳朵"：每天午餐/晚餐时播放英语课文录音，不求听懂每个词，关键是让大脑习惯英语的语音语调。', indent=True)
add_paragraph_styled(doc, '⑤ 朗读必不可少：每天早晨大声朗读20分钟，出声读才能形成语感。默读的英语是"哑巴英语"，考试听力一放就慌。', indent=True)

# --- 语文 ---
add_heading_styled(doc, '📚 语文（每日1.5小时）', level=2)
add_paragraph_styled(doc, '核心问题：基础尚可（68分），阅读理解缺乏答题方法，作文素材积累不足。', bold=True, indent=True)

create_table(doc,
    ['模块', '内容', '频次', '说明'],
    [
        ['古诗文', '七上+七下所有古诗文背诵默写', '每天20分钟', '逐篇过关，不能有错别字（默写错1个字=扣1分）'],
        ['文言文', '《世说新语》《论语》《狼》《卖油翁》等字词解释+翻译', '每周3次', '重点实词、虚词用法要整理成表'],
        ['现代文阅读', '每天做1篇记叙文阅读理解', '每天1篇', '做完对照答案总结答题模板，归纳答题套路'],
        ['写作', '每周写1篇600字作文', '每周1篇', '覆盖中考常考类型：亲情、成长、校园、自然'],
        ['名著阅读', '《朝花夕拾》《西游记》（初二上必读）', '每天30分钟', '课外时间完成，做简要读书笔记'],
    ],
    col_widths=[2.5, 5, 2, 5]
)

doc.add_paragraph()
add_paragraph_styled(doc, '💡 阅读答题模板（必背）：', bold=True)
add_paragraph_styled(doc, '• 概括题：谁 + 做了什么 + 结果', indent=True)
add_paragraph_styled(doc, '• 赏析题：修辞手法 + 内容 + 表达效果 + 作者情感', indent=True)
add_paragraph_styled(doc, '• 含义题：表层含义 + 深层含义（象征 / 比喻义）', indent=True)

# --- 历史 ---
add_heading_styled(doc, '🏛️ 历史 / 🗺️ 地理 / ⚖️ 政治（每日轮换一门，每科每周2次）', level=2)
add_paragraph_styled(doc, '这三科的特点是"背了就有分"，暑假集中突破性价比极高。从14-26分提到50-55分，靠的就是把七上七下课本完整过一遍。', bold=True, indent=True)

doc.add_paragraph()
add_paragraph_styled(doc, '🏛️ 历史学习方法：', bold=True)
add_paragraph_styled(doc, '• 核心方法：画时间轴。把七上（中国古代史：夏商周→秦汉→三国两晋南北朝→隋唐）和七下（宋元明清）的朝代顺序串成一条线。', indent=True)
add_paragraph_styled(doc, '• 课本目录是最好的复习提纲，先把目录背熟，再往里面填具体内容。', indent=True)
add_paragraph_styled(doc, '• 每日任务：每天背1课"知识点清单"，闭眼能完整复述才算过关。重点记忆：每个朝代的建立者、都城、重大事件、标志性人物。', indent=True)

doc.add_paragraph()
add_paragraph_styled(doc, '🗺️ 地理学习方法：', bold=True)
add_paragraph_styled(doc, '• 核心方法：地图是地理的灵魂。准备一张世界地图和中国地图贴在墙上，每天看。', indent=True)
add_paragraph_styled(doc, '• 七上重点：地球与地图（经纬度、自转公转）、天气与气候（气温降水图判读）。', indent=True)
add_paragraph_styled(doc, '• 七下重点：世界区域地理（亚洲、欧洲、非洲、美洲各区域的位置/地形/气候/经济）。', indent=True)
add_paragraph_styled(doc, '• 每日任务：每天记1个区域的地理特征，对着空白地图默写地理位置和地形区名称。', indent=True)

doc.add_paragraph()
add_paragraph_styled(doc, '⚖️ 政治学习方法：', bold=True)
add_paragraph_styled(doc, '• 核心方法：政治答题有固定套路——"是什么 + 为什么 + 怎么做"三步法。', indent=True)
add_paragraph_styled(doc, '• 七上重点：成长的节拍、友谊的天空、师长情谊、生命的思考。', indent=True)
add_paragraph_styled(doc, '• 七下重点：青春时光、做情绪的主人、在集体中成长、走进法治天地。', indent=True)
add_paragraph_styled(doc, '• 每日任务：每天背1课核心知识点 + 做2道材料分析题，模仿标准答案的答题格式。', indent=True)

# --- 生物 ---
add_heading_styled(doc, '🧬 生物（每日1小时）', level=2)
add_paragraph_styled(doc, '核心问题：有一定基础（47分），但概念理解不够深入，识图能力弱。', bold=True, indent=True)

create_table(doc,
    ['阶段', '时间', '内容', '学习方法'],
    [
        ['第1-2周', '七上复习', '细胞结构、生物圈、生态系统', '画细胞结构图、食物链/食物网示意图'],
        ['第3-5周', '七下复习', '人体系统（消化、呼吸、循环、泌尿、神经）', '每个系统画一张"结构→功能"流程图'],
        ['第6-8周', '初二上预习', '动物的主要类群（无脊椎→脊椎动物）', '对照课本做分类表格，理清进化脉络'],
    ],
    col_widths=[2, 2.5, 4.5, 5.5]
)

add_paragraph_styled(doc, '💡 生物学习的诀窍是"画图理解"而不是死记硬背。自己动手画一遍细胞图、人体系统示意图，比读十遍书都管用。', indent=True)

# ==================== 五、每周检查机制 ====================
add_heading_styled(doc, '五、每周检查机制', level=1)

add_paragraph_styled(doc, '每周日为检查日，家长需配合完成以下检查项目：', indent=True)

create_table(doc,
    ['检查项目', '内容', '方式', '合格标准'],
    [
        ['单词听写', '本周所背全部单词（约105个）', '家长念中文，孩子写英文', '正确率 ≥ 80%'],
        ['数学周测', '本周复习章节的基础题20道', '限时40分钟独立完成', '正确率 ≥ 70%'],
        ['古诗文默写', '本周背诵的古诗文篇目', '家长指定段落默写', '无错别字'],
        ['错题本检查', '本周数学+英语所有错题', '逐题检查是否按要求整理', '每题有错因+正解+同类题'],
        ['小科抽背', '本周轮换的小科背诵内容', '家长随机抽5个知识点提问', '能说出核心要点'],
    ],
    col_widths=[2.5, 4, 3.5, 4]
)

doc.add_paragraph()
add_paragraph_styled(doc, '📋 每周完成情况记录表（建议打印后贴在书桌前）：', bold=True)

create_table(doc,
    ['第N周', '日期', '数学\n(10分)', '英语\n(10分)', '语文\n(10分)', '生物\n(10分)', '小科\n(10分)', '总分\n(50分)', '家长签字'],
    [
        ['第1周', '___/___', '', '', '', '', '', '', ''],
        ['第2周', '___/___', '', '', '', '', '', '', ''],
        ['第3周', '___/___', '', '', '', '', '', '', ''],
        ['第4周', '___/___', '', '', '', '', '', '', ''],
        ['第5周', '___/___', '', '', '', '', '', '', ''],
        ['第6周', '___/___', '', '', '', '', '', '', ''],
        ['第7周', '___/___', '', '', '', '', '', '', ''],
        ['第8周', '___/___', '', '', '', '', '', '', ''],
    ],
    col_widths=[1.5, 2, 1.5, 1.5, 1.5, 1.5, 1.5, 2, 2.5]
)

# ==================== 六、给家长的建议 ====================
add_heading_styled(doc, '六、给家长的建议', level=1)

suggestions = [
    ('创造学习环境', '固定学习位置，学习时间全家保持安静。有一个固定的书桌比什么都重要。'),
    ('远离手机', '学习时段手机交给家长保管。手机是暑假学习最大的敌人——刷一条短视频，半个小时就没了。'),
    ('多鼓励少批评', '孩子目前基础薄弱，提升需要时间。每一点进步都值得表扬，批评只会加重厌学情绪。'),
    ('每日检查签字', '每天在孩子的学习任务完成表上签字。不需要辅导学习内容，只需要让孩子知道——爸妈在看，不能糊弄。'),
    ('关注心理健康', '成绩差距大容易产生畏难和厌学情绪，家长要及时沟通疏导。每天吃晚饭时可以聊聊今天学了什么，轻松的氛围下了解孩子的状态。'),
    ('适当奖励', '每完成一周的计划，安排一次孩子喜欢的活动：看电影、打篮球、和同学玩半天。奖励让坚持更有动力。'),
]

for i, (title_text, desc) in enumerate(suggestions, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title_text}：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = p.add_run(desc)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ==================== 七、关键提醒 ====================
add_heading_styled(doc, '七、关键提醒', level=1)

reminders = [
    ('🔑 数学和英语是翻盘的根本', '暑假8周，每天给数学和英语各安排近3小时。这两科一旦补上来，总分直接提升60分以上，而且为初二物理、化学打下基础。'),
    ('📖 小科背诵性价比最高', '历史地理政治从14-26分提到50-55分，靠的就是暑假把课本背一遍。这个分数提升对总分的贡献非常可观，而且不需要高智商，只需要下功夫。'),
    ('✍️ 错题本＞做新题', '错题本不是"抄题本"。每道错题都要分析：为什么错？正确思路是什么？下次遇到类似题怎么做？每周翻一次错题本，比做三套新卷子都有用。'),
    ('🎯 目标要踏实', '234.5到420+，看起来要提升185分很吓人。但分解到每科，不过是要把得分率从30%提到50-60%。只要把七年级的基础补扎实，这个目标完全能够做到。'),
    ('⏰ 坚持大于聪明', '这个计划最难的只有一件事——坚持。前两周可能很痛苦，但熬过去之后，学习节奏建立了，就会越来越顺。家长在这个阶段尤其要多鼓励。'),
]

for title_text, desc in reminders:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}：')
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = p.add_run(desc)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ==================== 结尾 ====================
doc.add_paragraph()
ending = doc.add_paragraph()
ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ending.add_run('计划是死的，执行是活的。每天坚持按作息表走，8周后一定会有质变。加油！🚀')
run.font.name = '楷体'
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(47, 84, 150)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

# ==================== 保存 ====================
output_path = r'D:\暑期提升计划\杨田安_暑期提升计划.docx'
doc.save(output_path)
print(f'Word doc saved to: {output_path}')
